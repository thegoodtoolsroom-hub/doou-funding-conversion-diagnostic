from __future__ import annotations

import json
import shutil
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document

from engine.commercial_qa import render_batch_quality_report
from engine.html_report_renderer import render_html_report, HTMLReportRendererError
from engine.limesurvey_decoder import DecodedResponse
from engine.pdf_report_renderer import render_pdf_report
from engine.report_data_builder import ReportDataBuilderError
from engine.scoring_engine import ScoredOrganisation


class BatchOutputError(Exception):
    """Raised when batch outputs cannot be produced safely."""


def generate_batch_outputs(
    decoded_responses: list[DecodedResponse],
    scored_organisations: list[ScoredOrganisation],
    report_bundle: list[dict[str, Any]],
    output_root: str | Path | None = None,
) -> list[dict[str, Any]]:
    if len(decoded_responses) != len(scored_organisations) or len(decoded_responses) != len(report_bundle):
        raise ReportDataBuilderError("Decoded responses, scored organisations and report data must align.")

    target_root = Path(output_root) if output_root is not None else Path("outputs")
    target_root.mkdir(parents=True, exist_ok=True)

    generated: list[dict[str, Any]] = []
    for decoded, scored, report_data in zip(decoded_responses, scored_organisations, report_bundle, strict=True):
        organisation_slug = _slugify(decoded.organisation_name)
        org_dir = target_root / organisation_slug
        org_dir.mkdir(parents=True, exist_ok=True)

        (org_dir / "report_data.json").write_text(json.dumps(report_data, indent=2), encoding="utf-8")

        response_audit = {
            "organisation_name": decoded.organisation_name,
            "decoded_scored_signal_count": decoded.decoded_scored_signals,
            "unknown_fields": decoded.unknown_fields,
            "missing_required_fields": decoded.missing_required_fields,
            "manual_review_flags": report_data.get("manual_review_flags", []),
        }
        (org_dir / "response_audit.json").write_text(json.dumps(response_audit, indent=2), encoding="utf-8")

        if report_data.get("manual_review_flags") and any("manual review" in flag.lower() for flag in report_data["manual_review_flags"]):
            (org_dir / "manual_review_note.txt").write_text(
                f"Manual review required for {decoded.organisation_name}\n",
                encoding="utf-8",
            )

        # Generate DOCX (placeholder format - kept for backward compatibility)
        template_path = _find_template_path()
        shutil.copy2(template_path, org_dir / "report.docx")
        _create_placeholder_docx(org_dir / "report.docx", report_data)

        # Generate HTML report
        html_path = org_dir / "report.html"
        html_render_result = {"status": "HTML_RENDER_ERROR", "html_path": None}
        try:
            html_content = render_html_report(report_data, html_path)
            html_render_result = {"status": "HTML_RENDERED", "html_path": str(html_path)}
        except HTMLReportRendererError as e:
            html_render_result = {"status": "HTML_RENDER_ERROR", "error": str(e), "html_path": None}

        # Generate PDF report from HTML
        pdf_path = org_dir / "report_branded.pdf"
        pdf_render_result = {"status": "PDF_RENDER_NOT_ATTEMPTED", "pdf_path": None}
        if html_render_result["status"] == "HTML_RENDERED":
            try:
                html_content = (org_dir / "report.html").read_text(encoding="utf-8")
                pdf_render_result = render_pdf_report(html_content, pdf_path)
            except Exception as e:
                pdf_render_result = {"status": "PDF_RENDER_ERROR", "error": str(e), "pdf_path": None}

        generated.append({
            "organisation_name": decoded.organisation_name,
            "organisation_slug": organisation_slug,
            "report_data_path": str(org_dir / "report_data.json"),
            "docx_path": str(org_dir / "report.docx"),
            "html_path": html_render_result.get("html_path"),
            "html_status": html_render_result.get("status"),
            "pdf_path": pdf_render_result.get("pdf_path") if pdf_render_result.get("status") == "PDF_RENDERED" else None,
            "pdf_status": pdf_render_result.get("status"),
        })

    _write_batch_index(target_root, generated)
    _write_batch_quality_report(target_root, generated)
    render_batch_quality_report(target_root)
    return generated


def _find_template_path() -> Path:
    candidates = [
        Path("templates/BATCH_03_REPORT_TEMPLATE_MASTER_UPLOAD_THIRD"),
        Path("templates"),
    ]
    for root in candidates:
        if root.exists():
            for path in root.glob("*.docx"):
                return path
    raise BatchOutputError("Could not find Word template .docx file.")


def _create_placeholder_docx(path: Path, report_data: dict[str, Any]) -> bool:
    doc = Document()
    doc.add_heading("DOO U Report", level=1)
    doc.add_paragraph(f"Organisation: {report_data.get('organisation_profile', {}).get('name', 'Unknown')}")
    doc.add_paragraph(f"Overall score: {report_data.get('overall_score', 'n/a')}")
    doc.add_paragraph("This placeholder report preserves the template and marks the current limitation of automated section injection.")
    doc.save(path)
    return True


def _slugify(value: str) -> str:
    slug = "".join(ch.lower() if ch.isalnum() else "-" for ch in value)
    slug = "-".join(part for part in slug.split("-") if part)
    return slug or "organisation"


def _write_batch_index(target_root: Path, generated: list[dict[str, Any]]) -> None:
    df = pd.DataFrame(generated)
    df.to_csv(target_root / "batch_index.csv", index=False)
    df.to_excel(target_root / "batch_index.xlsx", index=False)


def _write_batch_quality_report(target_root: Path, generated: list[dict[str, Any]]) -> None:
    rows = []
    for item in generated:
        rows.append(
            {
                "organisation": item["organisation_name"],
                "report_data": "✓" if item.get("report_data_path") else "✗",
                "html": item.get("html_status", "MISSING"),
                "pdf": item.get("pdf_status", "MISSING"),
                "docx": "✓" if item.get("docx_path") else "✗",
            }
        )
    
    html_rows = []
    for row in rows:
        html_rows.append(
            f"<tr>"
            f"<td>{row['organisation']}</td>"
            f"<td>{row['report_data']}</td>"
            f"<td>{row['html']}</td>"
            f"<td>{row['pdf']}</td>"
            f"<td>{row['docx']}</td>"
            f"</tr>"
        )
    
    html_table = (
        "<table border='1' style='border-collapse: collapse; width: 100%;'>"
        "<tr style='background-color: #20352D; color: white;'>"
        "<th style='padding: 8px;'>Organisation</th>"
        "<th style='padding: 8px;'>Report Data</th>"
        "<th style='padding: 8px;'>HTML Report</th>"
        "<th style='padding: 8px;'>PDF Report</th>"
        "<th style='padding: 8px;'>DOCX (Legacy)</th>"
        "</tr>"
        + "".join(html_rows)
        + "</table>"
    )
    
    html_content = f"""<html>
<head>
    <title>Batch Quality Report</title>
    <style>
        body {{ font-family: Arial, sans-serif; background-color: #F5F3F0; padding: 20px; }}
        h1 {{ color: #20352D; }}
        table {{ margin: 20px 0; background-color: white; }}
        th, td {{ padding: 8px; text-align: left; border: 1px solid #E0E0E0; }}
        th {{ background-color: #2E6B4F; color: white; }}
        tr:nth-child(even) {{ background-color: #F5F3F0; }}
    </style>
</head>
<body>
    <h1>Batch Quality Report</h1>
    <p>HTML/PDF Report Generation Status</p>
    {html_table}
</body>
</html>"""
    
    (target_root / "batch_quality_report.html").write_text(html_content, encoding="utf-8")
